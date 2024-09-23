import os
import subprocess
import datetime as dt
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from docx2pdf import convert
import docx


class InvoiceAutomation:

    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Invoice Automation")
        self.root.geometry("500x600")

        self.partner_label = tk.Label(self.root, text="Partner")
        self.partner_street_label = tk.Label(self.root, text="Partner Street")
        self.partner_zip_city_label = tk.Label(self.root, text="Partner ZIP Country")
        self.invoice_number_label = tk.Label(self.root, text="Invoice Number")
        self.service_description_label = tk.Label(self.root, text="Service Description")
        self.service_amount_label = tk.Label(self.root, text="Service Amount")
        self.service_single_price_label = tk.Label(self.root, text="Service Single Price")
        self.payment_method_label = tk.Label(self.root, text="Payment Method")


        self.payment_methods = {
            "Main Bank" : {
                'Recipient': 'Muhammad Kevin Fahlevi',
                'Bank': 'WISE',
                'IBAN': "HU87 7935 9021 8755 6002 1084",
                'BIC' : 'KHYDAOXXOIE'
            },
            "Second Bank" : {
                'Recipient': 'Muhammad Kevin Fahlevi',
                'Bank': 'REVOLUT',
                'IBAN': "HU65 9620 9733 1175 6565 9074",
                'BIC' : 'GIIUXYENYTR'
            },
            "Third Bank" : {
                'Recipient': 'Muhammad Kevin Fahlevi',
                'Bank': 'Citibank',
                'IBAN': "HU12 9754 7809 4312 4636 7694",
                'BIC' : 'HUTRKAAEEVM'
            }

        }

        self.partner_entry = tk.Entry(self.root)
        self.partner_street_entry = tk.Entry(self.root)
        self.partner_zip_city_country_entry = tk.Entry(self.root)
        self.invoice_number_entry = tk.Entry(self.root)
        self.service_description_entry = tk.Entry(self.root)
        self.service_amount_entry = tk.Entry(self.root)
        self.service_single_price_entry = tk.Entry(self.root)
        self.payment_method_entry = tk.Entry(self.root)

        self.payment_method = tk.StringVar(self.root)
        self.payment_method.set("Main Bank")

        self.payment_method_dropdown = tk.OptionMenu(self.root,self.payment_method, "Main Bank","Second Bank", "Private Bank" )
        self.create_button = tk.Button(self.root, text="Create Invoice", command=self.create_invoice)
        padding_options = {'fill': "x", "expand": True,"padx":5,"pady": 2}

        
        self.partner_label.pack(padding_options)
        self.partner_entry.pack(padding_options)
        self.partner_street_label.pack(padding_options)
        self.partner_street_entry.pack(padding_options)
        self.partner_zip_city_label.pack(padding_options)
        self.partner_zip_city_country_entry.pack(padding_options)
        self.invoice_number_label.pack(padding_options)
        self.invoice_number_entry.pack(padding_options)
        self.service_description_label.pack(padding_options)
        self.service_description_entry.pack(padding_options)
        self.service_amount_label.pack(padding_options)
        self.service_amount_entry.pack(padding_options)
        self.service_single_price_label.pack(padding_options)
        self.service_single_price_entry.pack(padding_options)
        self.payment_method_label.pack(padding_options)
        self.payment_method_dropdown.pack(padding_options)
        self.create_button.pack(padding_options)

        self.root.mainloop()

    @staticmethod
    def replace_text (paragraph, old_text, new_text):
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

    
    def create_invoice(self):
        doc = docx.Document('Current/template.docx')
        

        selected_payment_method = self.payment_methods[self.payment_method.get()]

        print(selected_payment_method['Bank'])

        try:
            replacement = {
                "[Date]" : dt.datetime.today().strftime('%Y-%m-%d'),
                "[Partner]" : self.partner_entry.get(),
                "[Partner Street]" : self.partner_street_entry.get(),
                "[Partner ZIP_City_Country]": self.partner_zip_city_country_entry.get(),
                "[Invoice Number]" : self.invoice_number_entry.get(),
                "[Service Description]" : self.service_description_entry.get(),
                "[Amount]" : self.service_amount_entry.get(),
                "[Single Price]" : f"${float(self.service_single_price_entry.get()):.2f}",
                "[Full Price]" : f'${float(self.service_amount_entry.get()) * float(self.service_single_price_entry.get()):.2f}',
                "[Recipient]" : selected_payment_method['Recipient'],
                "[Bank]" : selected_payment_method['Bank'],
                "[IBAN]" : selected_payment_method['IBAN'],
                "[BIC]" : selected_payment_method['BIC']
            }
        except ValueError:
            messagebox.showerror('Error','Invalid amount or price!')
            return
        
        for paragraph in list(doc.paragraphs):
            for old_text, new_text in replacement.items():
                self.replace_text(paragraph, old_text, new_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacement.items():
                            self.replace_text(paragraph, old_text, new_text)

    
        save_path = filedialog.asksaveasfilename(defaultextension='.pdf', filetypes= [('PDF documents', '*.pdf')])
        doc.save('filled.docx')
        convert('filled.docx', save_path)


        # subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', 'filled.docx', '--outdir', '.'], check=True)
        # os.rename('filled.pdf', save_path)

        messagebox.showinfo('Success', 'Invoice created and saved successfully!')

        #docx2pdf
        #from docx2pdf import convert
        #convert('filled.docx, save_path)


if __name__ == '__main__':
    InvoiceAutomation()
